#!/usr/bin/env python3
"""
CV Export Tool
Exports JSON Resume to PDF and DOCX formats with ATS-friendly template.

Usage:
    python export_cv.py                           # Default: all formats, pt-BR
    python export_cv.py --lang en-US              # English version
    python export_cv.py --format pdf              # PDF only
    python export_cv.py --format docx             # DOCX only
    python export_cv.py --all-langs               # All languages
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader

# Optional imports with graceful degradation
# Try xhtml2pdf first (pure Python, works everywhere)
try:
    from xhtml2pdf import pisa
    PDF_ENGINE = "xhtml2pdf"
except ImportError:
    # Fallback to weasyprint (requires GTK3 on Windows)
    try:
        from weasyprint import HTML as WeasyHTML
        PDF_ENGINE = "weasyprint"
    except (ImportError, OSError):
        PDF_ENGINE = None
        print("Warning: No PDF engine available. Install xhtml2pdf: pip install xhtml2pdf")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export disabled.")


# =============================================================================
# Configuration
# =============================================================================

BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
CV_FILE = BASE_DIR.parent / "cv_ed_costa.json"

# Labels for i18n
LABELS = {
    "pt-BR": {
        "summary": "Resumo Profissional",
        "experience": "Experiência Profissional",
        "education": "Formação Acadêmica",
        "skills": "Competências",
        "languages": "Idiomas",
        "certifications": "Certificações",
        "awards": "Premiações",
        "publications": "Publicações",
        "contact": "Contato",
        "present": "Atual",
    },
    "en-US": {
        "summary": "Professional Summary",
        "experience": "Professional Experience",
        "education": "Education",
        "skills": "Skills",
        "languages": "Languages",
        "certifications": "Certifications",
        "awards": "Awards",
        "publications": "Publications",
        "contact": "Contact",
        "present": "Present",
    }
}

# Language proficiency mapping for visual bars
LANGUAGE_LEVELS = {
    "Native": 100,
    "C2": 100,
    "C1": 90,
    "Fluent": 85,
    "Advanced": 80,
    "B2": 75,
    "B1": 60,
    "Intermediate": 55,
    "A2": 40,
    "Basic": 35,
    "A1": 25,
    "Elementary": 20,
}


# =============================================================================
# Jinja2 Filters
# =============================================================================

def format_date(date_str: str | None, lang: str = "pt-BR") -> str:
    """Format date string to readable format."""
    if not date_str:
        return ""

    try:
        date_obj = datetime.strptime(date_str[:10], "%Y-%m-%d")
        if lang == "pt-BR":
            months = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
                     "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
            return f"{months[date_obj.month - 1]}/{date_obj.year}"
        else:
            return date_obj.strftime("%b %Y")
    except (ValueError, TypeError):
        return str(date_str)[:7] if date_str else ""


def language_percent(lang_item: dict) -> int:
    """Convert language level to percentage for visual bars."""
    fluency = lang_item.get("fluency", "")
    level = lang_item.get("level", "")

    # Check level first (more specific)
    if level in LANGUAGE_LEVELS:
        return LANGUAGE_LEVELS[level]

    # Then check fluency
    for key, value in LANGUAGE_LEVELS.items():
        if key.lower() in fluency.lower():
            return value

    return 50  # Default


# =============================================================================
# Data Processing
# =============================================================================

def load_cv_data(cv_path: Path) -> dict:
    """Load and parse CV JSON file."""
    with open(cv_path, "r", encoding="utf-8") as f:
        return json.load(f)


def apply_i18n(data: dict, lang: str) -> dict:
    """Apply internationalization to CV data."""
    if lang == "pt-BR":
        return data  # Original is in pt-BR

    # Get i18n translations
    i18n = data.get("x-i18n", {}).get(lang, {})

    if not i18n:
        return data

    # Deep copy and apply translations
    result = json.loads(json.dumps(data))  # Deep copy

    # Apply basics translations
    if "basics" in i18n:
        result["basics"].update(i18n["basics"])

    # Apply work translations
    if "work" in i18n:
        for job in result.get("work", []):
            job_id = job.get("id", "")
            if job_id in i18n["work"]:
                job.update(i18n["work"][job_id])

    # Apply education translations
    if "education" in i18n:
        for edu in result.get("education", []):
            edu_id = edu.get("id", "")
            if edu_id in i18n["education"]:
                edu.update(i18n["education"][edu_id])

    # Apply awards translations
    if "awards" in i18n:
        for award in result.get("awards", []):
            award_id = award.get("id", "")
            if award_id in i18n["awards"]:
                award.update(i18n["awards"][award_id])

    # Apply certificates translations
    if "certificates" in i18n:
        for cert in result.get("certificates", []):
            cert_id = cert.get("id", "")
            if cert_id in i18n["certificates"]:
                cert.update(i18n["certificates"][cert_id])

    return result


def extract_ats_keywords(data: dict) -> list[str]:
    """Extract ATS-relevant keywords from CV data."""
    keywords = set()

    # From x-atsData
    ats_data = data.get("x-atsData", {}).get("keywords", {})
    for category in ats_data.values():
        if isinstance(category, list):
            keywords.update(category)

    # From skills
    for skill in data.get("skills", []):
        keywords.update(skill.get("keywords", []))

    # From work keywords
    for job in data.get("work", []):
        keywords.update(job.get("keywords", []))

    return sorted(keywords)[:50]  # Limit to top 50


def generate_output_filename(base_name: str, extension: str, output_dir: Path) -> Path:
    """
    Generate output filename following the pattern: CV - Ed Costa (2026-01-22).ext
    If file exists, append (n) before extension: CV - Ed Costa (2026-01-22) (1).ext

    Args:
        base_name: Base filename without extension (e.g., "CV - Ed Costa (2026-01-22)")
        extension: File extension (e.g., "pdf", "docx")
        output_dir: Output directory path

    Returns:
        Path object with unique filename
    """
    # Ensure extension doesn't have a leading dot
    extension = extension.lstrip('.')

    # Try the base filename first
    output_path = output_dir / f"{base_name}.{extension}"

    if not output_path.exists():
        return output_path

    # If file exists, find the next available number
    counter = 1
    while True:
        output_path = output_dir / f"{base_name} ({counter}).{extension}"
        if not output_path.exists():
            return output_path
        counter += 1


# =============================================================================
# PDF Export (xhtml2pdf or WeasyPrint)
# =============================================================================

def export_pdf(data: dict, template_name: str, output_path: Path, lang: str):
    """Export CV to PDF using xhtml2pdf or WeasyPrint."""
    if not PDF_ENGINE:
        print("  Skipping PDF: no PDF engine available")
        return False

    # Setup Jinja2
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))
    env.filters["format_date"] = lambda d: format_date(d, lang)
    env.filters["language_percent"] = language_percent

    # Load template
    template_file = f"{template_name}_template.html"
    template = env.get_template(template_file)

    # Prepare context
    context = {
        "lang": lang,
        "labels": LABELS[lang],
        "ats_keywords": extract_ats_keywords(data) if template_name == "ats" else None,
        **data
    }

    # Render HTML
    html_content = template.render(**context)

    # Generate PDF based on available engine
    if PDF_ENGINE == "xhtml2pdf":
        with open(output_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(html_content, dest=pdf_file, encoding='utf-8')
            if pisa_status.err:
                print(f"  Error creating PDF: {pisa_status.err}")
                return False
    else:  # weasyprint
        html = WeasyHTML(string=html_content, base_url=str(TEMPLATES_DIR))
        html.write_pdf(output_path)

    print(f"  Created: {output_path}")
    return True


# =============================================================================
# DOCX Export (python-docx)
# =============================================================================

def export_docx(data: dict, output_path: Path, lang: str):
    """Export CV to DOCX format (ATS-optimized)."""
    if not DOCX_AVAILABLE:
        print("  Skipping DOCX: python-docx not installed")
        return False

    labels = LABELS[lang]
    doc = Document()

    # Configure styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title style
    title_style = doc.styles.add_style('CVTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(44, 62, 80)

    # Section header style
    section_style = doc.styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Calibri'
    section_style.font.size = Pt(13)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(44, 62, 80)

    basics = data.get("basics", {})

    # Header - Name
    p = doc.add_paragraph(basics.get("name", ""), style='CVTitle')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header - Title
    p = doc.add_paragraph(basics.get("label", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(52, 73, 94)

    # Contact Info
    location = basics.get("location", {})
    contact_parts = [
        basics.get("email", ""),
        basics.get("phone", ""),
        f"{location.get('city', '')}, {location.get('region', '')}",
    ]
    if basics.get("url"):
        contact_parts.append(basics.get("url"))
    for profile in basics.get("profiles", []):
        contact_parts.append(profile.get("url", ""))

    p = doc.add_paragraph(" | ".join(filter(None, contact_parts)))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()  # Spacing

    # Summary
    doc.add_paragraph(labels["summary"], style='SectionHeader')
    doc.add_paragraph(basics.get("summary", ""))
    doc.add_paragraph()

    # Experience
    doc.add_paragraph(labels["experience"], style='SectionHeader')
    for job in data.get("work", []):
        # Position and date
        p = doc.add_paragraph()
        run = p.add_run(job.get("position", ""))
        run.bold = True

        date_str = f" | {format_date(job.get('startDate'), lang)} - {format_date(job.get('endDate'), lang) or labels['present']}"
        p.add_run(date_str)

        # Company
        p = doc.add_paragraph()
        run = p.add_run(job.get("name", ""))
        run.italic = True
        if job.get("location"):
            p.add_run(f" | {job.get('location')}")

        # Summary
        if job.get("summary"):
            doc.add_paragraph(job.get("summary"))

        # Highlights
        for highlight in job.get("highlights", []):
            p = doc.add_paragraph(highlight, style='List Bullet')

        doc.add_paragraph()  # Spacing

    # Education
    doc.add_paragraph(labels["education"], style='SectionHeader')
    for edu in data.get("education", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{edu.get('studyType', '')} - {edu.get('area', '')}")
        run.bold = True

        p = doc.add_paragraph()
        run = p.add_run(edu.get("institution", ""))
        run.italic = True

        date_str = f" | {format_date(edu.get('startDate'), lang)} - {format_date(edu.get('endDate'), lang) or labels['present']}"
        p.add_run(date_str)
        doc.add_paragraph()

    # Skills
    doc.add_paragraph(labels["skills"], style='SectionHeader')
    for skill in data.get("skills", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{skill.get('name', '')}: ")
        run.bold = True
        p.add_run(", ".join(skill.get("keywords", [])))
    doc.add_paragraph()

    # Languages
    doc.add_paragraph(labels["languages"], style='SectionHeader')
    for lang_item in data.get("languages", []):
        level_info = f" ({lang_item.get('level', '')})" if lang_item.get('level') else ""
        doc.add_paragraph(f"{lang_item.get('language', '')}: {lang_item.get('fluency', '')}{level_info}")
    doc.add_paragraph()

    # Certifications
    if data.get("certificates"):
        doc.add_paragraph(labels["certifications"], style='SectionHeader')
        for cert in data.get("certificates", []):
            status = f" [{cert.get('status')}]" if cert.get('status') else ""
            doc.add_paragraph(f"{cert.get('name', '')} - {cert.get('issuer', '')} ({format_date(cert.get('date'), lang)}){status}")
        doc.add_paragraph()

    # Awards
    if data.get("awards"):
        doc.add_paragraph(labels["awards"], style='SectionHeader')
        for award in data.get("awards", []):
            p = doc.add_paragraph()
            run = p.add_run(award.get("title", ""))
            run.bold = True
            p.add_run(f" - {award.get('awarder', '')} ({format_date(award.get('date'), lang)})")
            if award.get("summary"):
                doc.add_paragraph(award.get("summary"))
        doc.add_paragraph()

    # Publications
    if data.get("publications"):
        doc.add_paragraph(labels["publications"], style='SectionHeader')
        for pub in data.get("publications", []):
            p = doc.add_paragraph()
            run = p.add_run(pub.get("name", ""))
            run.bold = True
            doc.add_paragraph(f"{pub.get('publisher', '')}, {format_date(pub.get('releaseDate'), lang)}")

    # Save
    doc.save(output_path)
    print(f"  Created: {output_path}")
    return True


# =============================================================================
# Main
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Export CV to PDF and DOCX formats")
    parser.add_argument("--lang", choices=["pt-BR", "en-US"], default="pt-BR",
                       help="Language for export (default: pt-BR)")
    parser.add_argument("--format", choices=["pdf", "docx", "all"], default="all",
                       help="Output format (default: all)")
    parser.add_argument("--cv", type=Path, default=CV_FILE,
                       help="Path to CV JSON file")
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR,
                       help="Output directory")
    parser.add_argument("--all-langs", action="store_true",
                       help="Export all languages")

    args = parser.parse_args()

    # Create output directory
    args.output.mkdir(parents=True, exist_ok=True)

    # Load CV data
    print(f"Loading CV from: {args.cv}")
    raw_data = load_cv_data(args.cv)

    # Determine languages to export
    languages = ["pt-BR", "en-US"] if args.all_langs else [args.lang]

    # Use only ATS template
    template = "ats"
    formats = ["pdf", "docx"] if args.format == "all" else [args.format]

    print(f"\nExporting CV...")
    print(f"  Languages: {', '.join(languages)}")
    print(f"  Template: {template}")
    print(f"  Formats: {', '.join(formats)}")
    print()

    # Get name and current date for filename
    candidate_name = raw_data.get("basics", {}).get("name", "CV").split()
    # Extract first name and last name
    first_name = candidate_name[0] if candidate_name else "CV"
    last_name = candidate_name[-1] if len(candidate_name) > 1 else ""
    short_name = f"{first_name} {last_name}".strip()
    current_date = datetime.now().strftime("%Y-%m-%d")

    for lang in languages:
        print(f"[{lang}]")

        # Apply i18n
        data = apply_i18n(raw_data, lang)

        # PDF export (ATS template only)
        if "pdf" in formats:
            # Create base filename: CV - Ed Costa (2026-01-22) - pt-BR
            base_name = f"CV - {short_name} ({current_date}) - {lang}"
            output_file = generate_output_filename(base_name, "pdf", args.output)
            export_pdf(data, template, output_file, lang)

        # DOCX export (ATS-style)
        if "docx" in formats:
            # Create base filename: CV - Ed Costa (2026-01-22) - pt-BR
            base_name = f"CV - {short_name} ({current_date}) - {lang}"
            output_file = generate_output_filename(base_name, "docx", args.output)
            export_docx(data, output_file, lang)

        print()

    print("Export complete!")
    print(f"Files saved to: {args.output.absolute()}")


if __name__ == "__main__":
    main()
